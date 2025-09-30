#!/usr/bin/env python3
"""
Markdown to Word Document Converter
Converts IMPLEMENTATION.md, LITERATURE_REVIEW.md, and NOTES.md to .docx format
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def setup_document_styles(doc):
    """Setup custom styles for the Word document"""
    
    # Heading 1 Style
    try:
        heading1 = doc.styles['Heading 1']
    except:
        heading1 = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    
    # Heading 2 Style
    try:
        heading2 = doc.styles['Heading 2']
    except:
        heading2 = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    
    # Code Style
    try:
        code_style = doc.styles['Code']
    except:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)

def parse_markdown_content(content):
    """Parse markdown content and convert to Word document elements"""
    lines = content.split('\n')
    elements = []
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                elements.append(('code_block', '\n'.join(code_block_content)))
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Headers
        if line.startswith('# '):
            elements.append(('heading1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('heading2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('heading3', line[4:].strip()))
        elif line.startswith('#### '):
            elements.append(('heading4', line[5:].strip()))
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append(('bullet', line.strip()[2:]))
        # Numbered lists
        elif re.match(r'^\d+\. ', line.strip()):
            elements.append(('numbered', re.sub(r'^\d+\. ', '', line.strip())))
        # Empty lines
        elif line.strip() == '':
            elements.append(('empty', ''))
        # Regular text
        else:
            elements.append(('text', line))
    
    return elements

def convert_markdown_to_docx(md_file_path, output_path):
    """Convert a markdown file to Word document"""
    print(f"Converting {md_file_path} to {output_path}")
    
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    setup_document_styles(doc)
    
    # Parse markdown content
    elements = parse_markdown_content(content)
    
    # Convert elements to Word document
    for element_type, element_content in elements:
        if element_type == 'heading1':
            # Remove emoji and formatting
            clean_content = re.sub(r'[📚🔧🎯🏛️🔄🏗️💻🚀📋]', '', element_content).strip()
            para = doc.add_heading(clean_content, level=1)
            
        elif element_type == 'heading2':
            clean_content = re.sub(r'[📚🔧🎯🏛️🔄🏗️💻🚀📋⚙️📊💰🔍📉🛠️]', '', element_content).strip()
            para = doc.add_heading(clean_content, level=2)
            
        elif element_type == 'heading3':
            clean_content = re.sub(r'[📚🔧🎯🏛️🔄🏗️💻🚀📋⚙️📊💰🔍📉🛠️]', '', element_content).strip()
            para = doc.add_heading(clean_content, level=3)
            
        elif element_type == 'heading4':
            clean_content = re.sub(r'[📚🔧🎯🏛️🔄🏗️💻🚀📋⚙️📊💰🔍📉🛠️]', '', element_content).strip()
            para = doc.add_heading(clean_content, level=4)
            
        elif element_type == 'code_block':
            para = doc.add_paragraph(element_content)
            para.style = 'Code'
            
        elif element_type == 'bullet':
            clean_content = re.sub(r'[•▶◀●○]', '', element_content).strip()
            para = doc.add_paragraph(clean_content, style='List Bullet')
            
        elif element_type == 'numbered':
            para = doc.add_paragraph(element_content, style='List Number')
            
        elif element_type == 'text':
            if element_content.strip():
                # Remove markdown formatting
                clean_content = element_content
                clean_content = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_content)  # Bold
                clean_content = re.sub(r'\*(.*?)\*', r'\1', clean_content)      # Italic
                clean_content = re.sub(r'`(.*?)`', r'\1', clean_content)        # Code
                clean_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_content)  # Links
                
                para = doc.add_paragraph(clean_content)
        
        elif element_type == 'empty':
            doc.add_paragraph('')
    
    # Save document
    doc.save(output_path)
    print(f"✅ Successfully converted to {output_path}")

def main():
    """Main function to convert all markdown files"""
    # Define file mappings
    files_to_convert = [
        ('IMPLEMENTATION.md', 'IMPLEMENTATION.docx'),
        ('LITERATURE_REVIEW.md', 'LITERATURE_REVIEW.docx'),
        ('NOTES.md', 'NOTES.docx')
    ]
    
    print("🔄 Starting markdown to Word conversion process...")
    print("=" * 60)
    
    # Convert each file
    for md_file, docx_file in files_to_convert:
        if os.path.exists(md_file):
            try:
                convert_markdown_to_docx(md_file, docx_file)
            except Exception as e:
                print(f"❌ Error converting {md_file}: {str(e)}")
        else:
            print(f"⚠️ File not found: {md_file}")
    
    print("=" * 60)
    print("✅ Conversion process completed!")
    print("\n📁 Generated Word documents:")
    for _, docx_file in files_to_convert:
        if os.path.exists(docx_file):
            file_size = os.path.getsize(docx_file) / 1024  # KB
            print(f"  • {docx_file} ({file_size:.1f} KB)")

if __name__ == "__main__":
    main()